import re
import sys
import os
import pdfplumber
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from constants import (
    REQUIRED_SECTIONS, ATS_KEYWORDS, ACTION_VERBS,
    EMAIL_PATTERN, PHONE_PATTERN, SPECIAL_CHARS_PATTERN,
)
from parser import (
    parse_experience, parse_education, parse_projects,
    parse_certifications, parse_skills,
)


def main():
    from gui import launch
    launch()


def extract_text(filepath):
    if not os.path.isfile(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")

    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".txt":
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()

    elif ext == ".pdf":
        with pdfplumber.open(filepath) as pdf:
            pages = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
            return "\n".join(pages)

    elif ext == ".docx":
        with open(filepath, "rb") as f:
            doc = Document(f)
            return "\n".join(p.text for p in doc.paragraphs)

    else:
        raise ValueError(f"Unsupported file type: '{ext}'. Use (.txt / .pdf / .docx)")


def check_sections(text):
    results = {}
    lowered = text.lower()
    for section_name, pattern in REQUIRED_SECTIONS.items():
        results[section_name] = bool(re.search(pattern, lowered))
    return results


def score_resume(text):
    feedback = []
    word_count = len(text.split())

    sections = check_sections(text)
    sections_found = sum(sections.values())
    sections_score = int((sections_found / len(REQUIRED_SECTIONS)) * 30)
    missing_sections = [name.title() for name, found in sections.items() if not found]
    if missing_sections:
        feedback.append(f"Missing sections: {', '.join(missing_sections)}")
    else:
        feedback.append("All key sections found")

    keywords_found = [k for k in ATS_KEYWORDS if k in text.lower()]
    keywords_score = int(min(len(keywords_found) / 8, 1.0) * 25)
    feedback.append(f"{len(keywords_found)} ATS keywords found")

    found_verbs = [v for v in ACTION_VERBS if v in text.lower()]
    verbs_score = int(min(len(found_verbs) / 6, 1.0) * 20)
    feedback.append(f"{len(found_verbs)} action verbs found")

    if 300 <= word_count <= 800:
        word_count_score = 15
    elif 200 <= word_count < 300 or 800 < word_count <= 1000:
        word_count_score = 10
    else:
        word_count_score = 5
    feedback.append(f"Word count: {word_count} (best: 300-800)")

    formatting_score = 0

    if re.search(EMAIL_PATTERN, text):
        formatting_score += 4
        feedback.append("Email found")
    else:
        feedback.append("No email found")

    if re.search(PHONE_PATTERN, text):
        formatting_score += 3
        feedback.append("Phone number found")
    else:
        feedback.append("No phone number found")

    if not re.findall(SPECIAL_CHARS_PATTERN, text):
        formatting_score += 3
        feedback.append("Clean formatting")
    else:
        feedback.append("Special symbols found that may confuse ATS")

    total_score = (
        sections_score + keywords_score + verbs_score + word_count_score + formatting_score
    )

    return {
        "total_score": total_score,
        "breakdown": {
            "sections": sections_score,
            "keywords": keywords_score,
            "action_verbs": verbs_score,
            "word_count": word_count_score,
            "formatting": formatting_score,
        },
        "feedback": feedback,
    }


def build_resume(data, output_dir="output"):
    os.makedirs(output_dir, exist_ok=True)

    name = data.get("name", "").strip()
    email = data.get("email", "").strip()
    phone = data.get("phone", "").strip()
    linkedin = data.get("linkedin", "").strip()
    github = data.get("github", "").strip()
    summary = data.get("summary", "").strip()
    skills_raw = data.get("skills", "").strip()
    experience_raw = data.get("experience", "").strip()
    education_raw = data.get("education", "").strip()
    projects_raw = data.get("projects", "").strip()
    certifications_raw = data.get("certifications", "").strip()

    jobs = parse_experience(experience_raw)
    edu_entries = parse_education(education_raw)
    projects = parse_projects(projects_raw)
    certs = parse_certifications(certifications_raw)
    skill_lines = parse_skills(skills_raw)

    all_content_text = (summary + " " + experience_raw + " " + skills_raw).lower()
    missing_keywords = [
        k for k in ATS_KEYWORDS
        if k.lower() not in all_content_text
    ]
    if missing_keywords:
        injected = ", ".join(k.title() for k in missing_keywords[:6])
        skill_lines.append("Core Competencies: " + injected)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    def add_right_tab(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        tabs_elem = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "10224")
        tabs_elem.append(tab)
        pPr.append(tabs_elem)

    def add_section_heading(text):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(14)
        return paragraph

    def add_bullet(text):
        paragraph = doc.add_paragraph(style="List Paragraph")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        return paragraph

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(name)
    name_run.font.name = "Arial"
    name_run.font.size = Pt(26)

    contact_parts = [p for p in [email, phone] if p]
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.name = "Arial"
        contact_run.font.size = Pt(11)

    link_parts = [p for p in [linkedin, github] if p]
    if link_parts:
        link_para = doc.add_paragraph()
        link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        link_run = link_para.add_run(" | ".join(link_parts))
        link_run.font.name = "Arial"
        link_run.font.size = Pt(11)

    if summary:
        doc.add_paragraph()
        add_section_heading("SUMMARY")
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(summary)
        summary_run.font.name = "Arial"
        summary_run.font.size = Pt(11)

    if jobs:
        doc.add_paragraph()
        add_section_heading("WORK EXPERIENCE")

        for job in jobs:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(job["title"])
            title_run.bold = True
            title_run.font.name = "Arial"
            title_run.font.size = Pt(12)

            if job["company"] or job["date"]:
                company_para = doc.add_paragraph()
                add_right_tab(company_para)
                company_run = company_para.add_run(job["company"])
                company_run.bold = True
                company_run.font.name = "Arial"
                company_run.font.size = Pt(10)
                if job["date"]:
                    date_run = company_para.add_run("\t" + job["date"])
                    date_run.bold = True
                    date_run.font.name = "Arial"
                    date_run.font.size = Pt(10)

            for bullet in job["bullets"]:
                add_bullet(bullet)

    if edu_entries:
        doc.add_paragraph()
        add_section_heading("EDUCATION")

        for entry in edu_entries:
            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(entry["degree"])
            degree_run.font.name = "Arial"
            degree_run.font.size = Pt(12)

            if entry["specialization"]:
                spec_para = doc.add_paragraph()
                spec_run = spec_para.add_run(entry["specialization"])
                spec_run.font.name = "Arial"
                spec_run.font.size = Pt(12)

            if entry["institution"] or entry["details"]:
                inst_para = doc.add_paragraph()
                add_right_tab(inst_para)
                inst_run = inst_para.add_run(entry["institution"])
                inst_run.font.name = "Arial"
                inst_run.font.size = Pt(10)
                if entry["details"]:
                    details_run = inst_para.add_run("\t" + entry["details"])
                    details_run.font.name = "Arial"
                    details_run.font.size = Pt(10)

    if skill_lines:
        doc.add_paragraph()
        add_section_heading("SKILLS")
        for line in skill_lines:
            add_bullet(line)

    if projects:
        doc.add_paragraph()
        add_section_heading("PROJECTS")

        for project in projects:
            proj_para = doc.add_paragraph()
            proj_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_right_tab(proj_para)

            title_run = proj_para.add_run(project["title"])
            title_run.bold = True
            title_run.font.name = "Arial"
            title_run.font.size = Pt(12)

            if project["date"]:
                date_run = proj_para.add_run("\t" + project["date"])
                date_run.bold = True
                date_run.font.name = "Arial"
                date_run.font.size = Pt(10)

            for bullet in project["bullets"]:
                add_bullet(bullet)

    if certs:
        doc.add_paragraph()
        add_section_heading("CERTIFICATIONS")

        for cert in certs:
            cert_para = doc.add_paragraph(style="List Paragraph")
            cert_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            name_run = cert_para.add_run(cert["name"])
            name_run.font.name = "Arial"
            name_run.font.size = Pt(11)

            if cert["issuer"]:
                sep_run = cert_para.add_run(" · ")
                sep_run.bold = True
                sep_run.font.name = "Arial"
                sep_run.font.size = Pt(14)

                issuer_run = cert_para.add_run(cert["issuer"])
                issuer_run.font.name = "Arial"
                issuer_run.font.size = Pt(11)

    safe_name = re.sub(r"[^\w\s-]", "", name).strip().replace(" ", "_")
    filename = f"{safe_name}_resume.docx" if safe_name else "generated_resume.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    generated_text = extract_text(filepath)
    result = score_resume(generated_text)

    return {
        "filepath": filepath,
        "filename": filename,
        "score_result": result,
    }


if __name__ == "__main__":
    main()

